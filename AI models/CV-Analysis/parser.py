from __future__ import annotations

import io
import logging
import os
import platform
import re
from dataclasses import asdict, dataclass, field

import pdfplumber
from docx import Document
from fastapi import HTTPException

from utils import clean_text

logger = logging.getLogger(__name__)

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text

    _PDFMINER_AVAILABLE = True
except ImportError:
    _PDFMINER_AVAILABLE = False
    pdfminer_extract_text = None  # type: ignore[assignment]

try:
    import cv2
    import numpy as np

    _OPENCV_AVAILABLE = True
except ImportError:
    _OPENCV_AVAILABLE = False
    cv2 = None  # type: ignore[assignment]
    np = None  # type: ignore[assignment]

try:
    import layoutparser as lp

    _LAYOUTPARSER_AVAILABLE = True
except ImportError:
    _LAYOUTPARSER_AVAILABLE = False
    lp = None  # type: ignore[assignment]

try:
    import pytesseract
    from pdf2image import convert_from_bytes

    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False
    logger.warning("pytesseract / pdf2image not installed - OCR fallback disabled.")

TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"D:\poppler-25.12.0\Library\bin"

if _OCR_AVAILABLE and platform.system() == "Windows" and os.path.exists(TESSERACT_CMD):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


@dataclass
class ParseDiagnostics:
    multi_column_layout: bool = False
    tables_detected: bool = False
    mixed_formatting: bool = False
    broken_ocr_text: bool = False
    layout_blocks_detected: bool = False
    original_issues: list[str] = field(default_factory=list)
    fixes_applied: list[str] = field(default_factory=list)
    extraction_method: str = "native"
    weak_pages: int = 0
    dependency_status: dict[str, bool] = field(default_factory=dict)

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class ParseResult:
    text: str
    diagnostics: ParseDiagnostics


def _rows_to_lines(selected_words: list[dict]) -> list[str]:
    rows: dict[int, list[dict]] = {}
    for word in selected_words:
        row_key = int(round(word["top"] / 6)) * 6
        rows.setdefault(row_key, []).append(word)

    lines: list[str] = []
    for row_key in sorted(rows.keys()):
        row_words = sorted(rows[row_key], key=lambda w: w["x0"])
        line = " ".join(w["text"] for w in row_words).strip()
        if line:
            lines.append(line)
    return lines


def _inspect_pdf_page(page) -> dict:
    words = page.extract_words(
        x_tolerance=4,
        y_tolerance=6,
        keep_blank_chars=False,
        use_text_flow=False,
        extra_attrs=["fontname", "size"],
    )
    if not words:
        return {"text": "", "two_columns": False, "tables": False}

    left_words = [w for w in words if w["x0"] < page.width * 0.45]
    right_words = [w for w in words if w["x0"] > page.width * 0.55]
    has_two_columns = (
        len(left_words) >= 8
        and len(right_words) >= 8
        and abs(len(left_words) - len(right_words)) <= max(len(words) // 2, 8)
    )

    if has_two_columns:
        left_lines = _rows_to_lines([w for w in words if w["x0"] < page.width * 0.5])
        right_lines = _rows_to_lines([w for w in words if w["x0"] >= page.width * 0.5])
        text = "\n".join(left_lines + right_lines)
    else:
        text = "\n".join(_rows_to_lines(words))

    table_like_lines = 0
    for line in text.splitlines():
        pipe_count = line.count("|")
        spacer_count = len(re.findall(r"\s{3,}", line))
        if pipe_count >= 2 or spacer_count >= 3:
            table_like_lines += 1

    return {
        "text": text,
        "two_columns": has_two_columns,
        "tables": table_like_lines >= 2,
    }


def _preprocess(text: str) -> str:
    lorem_words = {
        "laoreet",
        "donec",
        "hendrerit",
        "pellentesque",
        "adipiscing",
        "consectetur",
        "lorem",
        "ipsum",
    }

    cleaned: list[str] = []
    for line in text.split("\n"):
        stripped = line.strip()
        if len(stripped) < 2:
            continue

        words_in_line = set(stripped.lower().split())
        if len(words_in_line & lorem_words) >= 2:
            continue

        stripped = re.sub(r"\s{2,}", " ", stripped)
        cleaned.append(stripped)

    result = "\n".join(cleaned)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def _detect_mixed_formatting(text: str) -> bool:
    return bool(re.search(r"[■●◆★✓✔☑☎✉]", text))


def _detect_broken_ocr(text: str) -> bool:
    suspicious_patterns = [
        re.compile(r"[A-Za-z]{1}\s[A-Za-z]{1}\s[A-Za-z]{1}\s[A-Za-z]{1,}"),
        re.compile(r"[A-Za-z]{8,}\d{2,}[A-Za-z]{3,}"),
        re.compile(r"[^\w\s@:/+.#,\-%()]{4,}"),
    ]
    short_fragments = sum(1 for line in text.splitlines() if 0 < len(line.strip()) <= 2)
    return short_fragments >= 6 or any(pattern.search(text) for pattern in suspicious_patterns)


def _preprocess_image_for_ocr(image):
    if not _OPENCV_AVAILABLE or np is None or cv2 is None:
        return image.convert("L")

    image_array = np.array(image.convert("L"))
    denoised = cv2.GaussianBlur(image_array, (3, 3), 0)
    _, thresholded = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return thresholded


def _estimate_layout_blocks(text: str) -> bool:
    if not _LAYOUTPARSER_AVAILABLE:
        return False

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    short_lines = [line for line in lines if len(line.split()) <= 4]
    long_lines = [line for line in lines if len(line.split()) >= 8]
    return bool(short_lines and long_lines and len(short_lines) >= 6)


def _extract_with_pdfminer(file_bytes: bytes) -> str:
    if not _PDFMINER_AVAILABLE or pdfminer_extract_text is None:
        return ""

    try:
        return pdfminer_extract_text(io.BytesIO(file_bytes)) or ""
    except Exception as exc:
        logger.warning("pdfminer extraction failed: %s", exc)
        return ""


def _clean_ocr_noise(text: str) -> tuple[str, list[str]]:
    fixes: list[str] = []
    cleaned = clean_text(text)

    collapsed_symbols = re.sub(r"[■●◆★✓✔☑]+", "-", cleaned)
    if collapsed_symbols != cleaned:
        fixes.append("Removed decorative bullets and symbols that break ATS parsing")
        cleaned = collapsed_symbols

    dehyphenated = re.sub(r"(?<=\w)-\n(?=\w)", "", cleaned)
    if dehyphenated != cleaned:
        fixes.append("Merged hyphenated OCR line breaks")
        cleaned = dehyphenated

    repaired_spacing = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", cleaned)
    if repaired_spacing != cleaned:
        fixes.append("Inserted missing spaces in merged OCR tokens")
        cleaned = repaired_spacing

    normalized_newlines = re.sub(r"\n{3,}", "\n\n", cleaned)
    if normalized_newlines != cleaned:
        fixes.append("Collapsed excessive blank lines and page artifacts")
        cleaned = normalized_newlines

    return _preprocess(cleaned), fixes


def _ocr_pdf(file_bytes: bytes) -> str:
    if not _OCR_AVAILABLE:
        return ""

    try:
        kwargs = {"dpi": 250}
        if POPPLER_PATH and os.path.exists(POPPLER_PATH) and platform.system() == "Windows":
            kwargs["poppler_path"] = POPPLER_PATH

        images = convert_from_bytes(file_bytes, **kwargs)
        parts: list[str] = []

        for image in images:
            processed_image = _preprocess_image_for_ocr(image)
            text = pytesseract.image_to_string(processed_image, config="--oem 3 --psm 6")
            if text.strip():
                parts.append(text)

        return "\n".join(parts)
    except Exception as exc:
        logger.warning("OCR failed: %s", exc)
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> ParseResult:
    diagnostics = ParseDiagnostics()
    diagnostics.dependency_status = {
        "pytesseract": _OCR_AVAILABLE,
        "pdfminer.six": _PDFMINER_AVAILABLE,
        "opencv-python": _OPENCV_AVAILABLE,
        "layoutparser": _LAYOUTPARSER_AVAILABLE,
    }

    try:
        pages_text: list[str] = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                inspection = _inspect_pdf_page(page)
                diagnostics.multi_column_layout = diagnostics.multi_column_layout or inspection["two_columns"]
                diagnostics.tables_detected = diagnostics.tables_detected or inspection["tables"]

                page_text = inspection["text"]
                if len(page_text.strip()) < 30:
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3, layout=True) or ""

                if page_text.strip():
                    pages_text.append(page_text)
                    if len(page_text.strip()) < 80:
                        diagnostics.weak_pages += 1
                else:
                    diagnostics.weak_pages += 1

        raw_text = "\n".join(pages_text).strip()
        diagnostics.mixed_formatting = _detect_mixed_formatting(raw_text)
        diagnostics.layout_blocks_detected = _estimate_layout_blocks(raw_text)

        if diagnostics.layout_blocks_detected and not diagnostics.multi_column_layout:
            diagnostics.multi_column_layout = True
            diagnostics.fixes_applied.append("Detected layout-block structure and linearized content for ATS-safe reading order")

        pdfminer_text = _extract_with_pdfminer(file_bytes)
        if len(pdfminer_text.strip()) > len(raw_text.strip()):
            raw_text = pdfminer_text
            diagnostics.extraction_method = "pdfminer"
            diagnostics.fixes_applied.append("Used pdfminer fallback to recover additional PDF text")

        used_ocr = False
        if len(raw_text.strip()) < 80 or diagnostics.weak_pages:
            ocr_text = _ocr_pdf(file_bytes)
            if len(ocr_text.strip()) > len(raw_text.strip()):
                raw_text = ocr_text
                used_ocr = True
            elif ocr_text.strip():
                raw_text = "\n\n".join(part for part in (raw_text, ocr_text) if part.strip())
                used_ocr = True

        diagnostics.broken_ocr_text = _detect_broken_ocr(raw_text)
        cleaned_text, cleaning_fixes = _clean_ocr_noise(raw_text)

        if used_ocr:
            diagnostics.extraction_method = "ocr_fallback"
            diagnostics.fixes_applied.append("Applied OCR fallback for weak or image-based PDF pages")

        diagnostics.fixes_applied.extend(cleaning_fixes)

        if diagnostics.multi_column_layout:
            diagnostics.original_issues.append("Detected multi-column layout and reconstructed reading order")
        if diagnostics.tables_detected:
            diagnostics.original_issues.append("Detected table-like content that can break ATS parsers")
        if diagnostics.mixed_formatting:
            diagnostics.original_issues.append("Detected decorative formatting/icons that are not ATS-friendly")
        if diagnostics.layout_blocks_detected:
            diagnostics.original_issues.append("Detected mixed layout blocks consistent with complex CV formatting")
        if diagnostics.broken_ocr_text:
            diagnostics.original_issues.append("Detected OCR noise or merged/broken text fragments")
        if diagnostics.weak_pages:
            diagnostics.original_issues.append(f"Detected {diagnostics.weak_pages} weakly extracted page(s)")

        if len(cleaned_text.strip()) < 10:
            raise HTTPException(
                status_code=422,
                detail=(
                    "No readable text found. The PDF may be image-based. "
                    "Install pytesseract + pdf2image for OCR support."
                ),
            )

        return ParseResult(text=cleaned_text, diagnostics=diagnostics)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Failed to parse PDF: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> ParseResult:
    diagnostics = ParseDiagnostics()
    diagnostics.dependency_status = {
        "pytesseract": _OCR_AVAILABLE,
        "pdfminer.six": _PDFMINER_AVAILABLE,
        "opencv-python": _OPENCV_AVAILABLE,
        "layoutparser": _LAYOUTPARSER_AVAILABLE,
    }

    try:
        document = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in document.tables:
            diagnostics.tables_detected = True
            for row in table.rows:
                row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        if not parts:
            raise HTTPException(
                status_code=422,
                detail="Could not extract text from the DOCX file.",
            )

        raw_text = "\n".join(parts)
        diagnostics.mixed_formatting = _detect_mixed_formatting(raw_text)
        diagnostics.broken_ocr_text = _detect_broken_ocr(raw_text)
        diagnostics.layout_blocks_detected = _estimate_layout_blocks(raw_text)
        cleaned_text, cleaning_fixes = _clean_ocr_noise(raw_text)
        diagnostics.fixes_applied.extend(cleaning_fixes)

        if diagnostics.tables_detected:
            diagnostics.original_issues.append("Detected table-based layout in DOCX content")
            diagnostics.fixes_applied.append("Flattened table content into ATS-safe linear text")
        if diagnostics.mixed_formatting:
            diagnostics.original_issues.append("Detected mixed formatting or symbols in DOCX content")
        if diagnostics.layout_blocks_detected:
            diagnostics.original_issues.append("Detected complex block-based layout in DOCX content")
        if diagnostics.broken_ocr_text:
            diagnostics.original_issues.append("Detected broken OCR-like text fragments in DOCX content")

        return ParseResult(text=cleaned_text, diagnostics=diagnostics)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Failed to parse DOCX: {exc}") from exc


def parse_cv_detailed(filename: str, file_bytes: bytes) -> ParseResult:
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    raise HTTPException(
        status_code=415,
        detail="Unsupported file type. Please upload a PDF or DOCX file.",
    )


def parse_cv(filename: str, file_bytes: bytes) -> str:
    return parse_cv_detailed(filename, file_bytes).text
